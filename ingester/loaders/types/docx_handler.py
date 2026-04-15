"""Docx handler for v3 Drive loader (session 18).

Extracts text, headings, tables, and inline images from .docx files using
python-docx. Follows the same handler protocol as pdf_handler and
google_doc_handler:

  1. Walk the document in order, producing an ordered stream of
     {kind: "text"|"heading"|"image"|"table"} entries.
  2. Stitch the stream into a single string with [SECTION N] markers
     at heading boundaries and [IMAGE #N: ocr_text] for inline images.
  3. Return an ExtractResult for the dispatcher.

Heading convention:
  Uses [SECTION N] markers (same as Google Docs) since docx documents
  have the same paragraph/heading structure. derive_locator() renders
  these as §N / §§N-M.

Table convention:
  Tables are serialized as pipe-delimited text rows and spliced inline
  at their document position. This preserves content from supplement
  lists, lab ranges, protocol schedules, etc.

Image convention:
  Inline images (InlineShape objects with an embedded image relationship)
  are extracted and sent through the shared GeminiVisionClient for OCR.
  The OCR text is spliced inline as [IMAGE #N: ocr_text], same as the
  Google Doc handler. Floating/anchor images are not accessible via
  python-docx and are skipped with a warning.

Public interface:
  extract_from_path(path, *, vision_client=None, use_cache=True)
      -> ExtractResult

  extract(drive_file, drive_client, config) -> ExtractResult
      -> dispatcher entrypoint
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

from ingester.loaders.types import (
    ExtractResult,
    make_page_marker,
)

log = logging.getLogger(__name__)

# python-docx heading style names start with "Heading" followed by
# a space and the level number (1-9). We detect these to emit
# [SECTION N] markers.
_HEADING_STYLE_PREFIX = "Heading"


# ----------------------------------------------------------------------------
# Document walk — produces an ordered stream
# ----------------------------------------------------------------------------

def _walk_document(doc) -> list[dict]:
    """Walk a python-docx Document in body order, producing a stream of
    {kind: "text"|"heading"|"image"|"table"} entries.

    python-docx's Document.element.body iterates over paragraphs and
    tables in document order via the underlying XML. We iterate the
    body's child elements directly to preserve ordering between
    paragraphs and tables (Document.paragraphs and Document.tables
    are separate flat lists that lose interleaving order).
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    stream: list[dict] = []
    body = doc.element.body

    for child in body:
        tag = child.tag

        # --- Paragraph (including headings) ---
        if tag == qn("w:p"):
            para = Paragraph(child, doc)
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            # Check for inline images in this paragraph
            inline_images = child.findall(
                ".//wp:inline//a:blip",
                {
                    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
                    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                },
            )

            # Detect heading
            is_heading = (
                style_name.startswith(_HEADING_STYLE_PREFIX)
                and len(style_name) > len(_HEADING_STYLE_PREFIX)
            )
            if is_heading:
                try:
                    level = int(style_name[len(_HEADING_STYLE_PREFIX):].strip())
                except ValueError:
                    level = 1
                stream.append({
                    "kind": "heading",
                    "level": level,
                    "text": text,
                })

            # Emit paragraph text (even for headings — mirrors google_doc
            # handler which emits heading text as a paragraph too)
            if text:
                stream.append({"kind": "text", "text": text})

            # Emit any inline images found in this paragraph
            for blip in inline_images:
                r_embed = blip.get(qn("r:embed"))
                if r_embed:
                    stream.append({
                        "kind": "image",
                        "rel_id": r_embed,
                    })

        # --- Table ---
        elif tag == qn("w:tbl"):
            tbl = Table(child, doc)
            rows_text: list[list[str]] = []
            for row in tbl.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(cells)
            if rows_text:
                stream.append({"kind": "table", "rows": rows_text})

    return stream


# ----------------------------------------------------------------------------
# Image extraction from docx relationships
# ----------------------------------------------------------------------------

def _extract_image_bytes(doc, rel_id: str) -> tuple[bytes, str]:
    """Extract image bytes from a docx relationship ID.

    Returns (image_bytes, mime_type).
    Raises RuntimeError if the relationship or image part is missing.
    """
    try:
        rel = doc.part.rels[rel_id]
    except KeyError:
        raise RuntimeError(f"Relationship {rel_id} not found in docx")

    image_part = rel.target_part
    img_bytes = image_part.blob
    content_type = image_part.content_type or "image/png"
    return img_bytes, content_type


# ----------------------------------------------------------------------------
# Table serialization
# ----------------------------------------------------------------------------

def _serialize_table(rows: list[list[str]]) -> str:
    """Serialize a table as pipe-delimited text rows.

    Produces markdown-style table with header separator. If the table
    has at least one row, the first row is treated as a header.
    """
    if not rows:
        return ""
    lines: list[str] = []
    for i, row in enumerate(rows):
        line = " | ".join(row)
        lines.append(line)
        if i == 0 and len(rows) > 1:
            # Header separator
            lines.append(" | ".join(["---"] * len(row)))
    return "\n".join(lines)


# ----------------------------------------------------------------------------
# Stream stitch + image OCR
# ----------------------------------------------------------------------------

def _stitch_stream(
    stream: list[dict],
    doc,
    vision_client,
    use_cache: bool,
) -> tuple[str, int, list[dict], int]:
    """Walk the ordered stream, OCR images, and produce a single stitched
    text string with [SECTION N] markers at headings and [IMAGE #N: ...]
    at image positions.

    Returns (stitched_text, image_count, per_image_records, section_count).
    Mirrors google_doc_handler.stitch_stream shape.
    """
    parts: list[str] = []
    image_count = 0
    section_count = 0
    per_image: list[dict] = []

    for entry in stream:
        kind = entry["kind"]

        if kind == "heading":
            section_count += 1
            marker = make_page_marker("SECTION", section_count)
            parts.append(marker)
            # Heading text is emitted by the subsequent "text" entry
            continue

        if kind == "text":
            txt = entry["text"]
            if txt.strip():
                parts.append(txt)
            continue

        if kind == "table":
            table_text = _serialize_table(entry["rows"])
            if table_text.strip():
                parts.append(table_text)
            continue

        if kind == "image":
            image_count += 1
            rel_id = entry["rel_id"]
            record = {"index": image_count, "rel_id": rel_id}

            # Extract image bytes from docx package
            try:
                img_bytes, img_mime = _extract_image_bytes(doc, rel_id)
                record["byte_size"] = len(img_bytes)
                record["mime_type"] = img_mime
            except Exception as e:
                reason = str(e)[:80]
                record["extract_failed"] = reason
                parts.append(
                    f"\n\n[IMAGE #{image_count}: EXTRACT_FAILED — {reason}]\n\n"
                )
                per_image.append(record)
                continue

            # OCR via shared vision client (if available)
            if vision_client is None:
                record["skipped"] = "no_vision_client"
                per_image.append(record)
                continue

            ocr = vision_client.ocr_image(img_bytes, img_mime, use_cache=use_cache)
            record["sha256"] = ocr.sha256
            record["is_decorative"] = ocr.is_decorative
            record["failed"] = ocr.failed
            record["ocr_text_preview"] = ocr.ocr_text[:200]
            record["vision_input_tokens"] = ocr.vision_input_tokens
            record["vision_output_tokens"] = ocr.vision_output_tokens

            if ocr.is_decorative:
                pass  # Drop from stream
            elif ocr.failed:
                parts.append(
                    f"\n\n[IMAGE #{image_count}: OCR_FAILED — {ocr.failure_reason}]\n\n"
                )
            else:
                parts.append(
                    f"\n\n[IMAGE #{image_count}: {ocr.ocr_text.strip()}]\n\n"
                )
            per_image.append(record)
            continue

    stitched = "\n\n".join(parts)
    return stitched, image_count, per_image, section_count


# ----------------------------------------------------------------------------
# extract_from_path — main extraction entrypoint
# ----------------------------------------------------------------------------

def extract_from_path(
    path: Path,
    *,
    vision_client=None,
    use_cache: bool = True,
) -> ExtractResult:
    """Extract text from a .docx file on local disk.

    Args:
      path: local filesystem path to the .docx file
      vision_client: optional pre-constructed GeminiVisionClient for
        inline image OCR. If None, images are skipped silently.
      use_cache: whether to use the shared OCR cache (default True).

    Returns:
      ExtractResult with stitched_text containing [SECTION N] markers
      at headings and [IMAGE #N: ...] at image positions.
    """
    from docx import Document

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Docx not found: {path}")

    doc = Document(str(path))
    stream = _walk_document(doc)

    # Count headings and images for diagnostics
    heading_count = sum(1 for e in stream if e["kind"] == "heading")
    image_count_pre = sum(1 for e in stream if e["kind"] == "image")
    table_count = sum(1 for e in stream if e["kind"] == "table")

    # Capture vision ledger before stitch (for delta calculation)
    images_seen_before = 0
    images_ocr_before = 0
    cost_before = 0.0
    if vision_client is not None:
        images_seen_before = vision_client.ledger.images_seen
        images_ocr_before = vision_client.ledger.images_ocr_called
        cost_before = vision_client.ledger.vision_cost_usd

    stitched, image_count, per_image, section_count = _stitch_stream(
        stream, doc, vision_client, use_cache
    )

    # Compute vision ledger deltas
    delta_seen = 0
    delta_called = 0
    delta_cost = 0.0
    if vision_client is not None:
        delta_seen = vision_client.ledger.images_seen - images_seen_before
        delta_called = vision_client.ledger.images_ocr_called - images_ocr_before
        delta_cost = vision_client.ledger.vision_cost_usd - cost_before

    warnings: list[str] = []
    if heading_count == 0:
        warnings.append(
            "no heading styles found — display_locator will be empty "
            "(BACKLOG #32: paragraph fallback)"
        )

    return ExtractResult(
        stitched_text=stitched,
        extraction_method="docx_python_docx",
        source_unit_label="section",
        pages_total=0,  # docx has no page concept at extraction time
        units_total=heading_count,
        images_seen=delta_seen,
        images_ocr_called=delta_called,
        vision_cost_usd=round(delta_cost, 6),
        warnings=warnings,
        extra={
            "per_image_records": per_image,
            "image_count_in_stream": image_count,
            "section_count": section_count,
            "table_count": table_count,
        },
    )


# ----------------------------------------------------------------------------
# Dispatcher entrypoint (called by drive_loader_v3._dispatch_file)
# ----------------------------------------------------------------------------

def extract(drive_file: dict, drive_client, config) -> ExtractResult:
    """Dispatcher entrypoint — downloads the .docx from Drive, delegates
    to extract_from_path(), returns the ExtractResult.

    `drive_file` is the Drive API metadata dict (id, name, mimeType, ...).
    `drive_client` is the v1/v2 DriveClient wrapper.
    `config` carries the shared `vision_client` and `use_cache` flag.
    Same shim shape as pdf_handler.extract and google_doc_handler.extract.
    """
    import tempfile

    file_id = drive_file["id"]
    file_name = drive_file.get("name", file_id)

    vision_client = getattr(config, "vision_client", None) if config else None
    use_cache = getattr(config, "use_cache", True) if config else True

    # Download binary from Drive
    data = drive_client.download_file_bytes(file_id)
    with tempfile.NamedTemporaryFile(
        suffix=".docx", delete=False, prefix="v3_docx_"
    ) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)

    try:
        result = extract_from_path(
            tmp_path,
            vision_client=vision_client,
            use_cache=use_cache,
        )
    finally:
        try:
            tmp_path.unlink()
        except OSError:
            pass

    # Tag extra with Drive-side metadata for diagnostics
    result.extra["drive_file_id"] = file_id
    result.extra["drive_file_name"] = file_name
    return result
