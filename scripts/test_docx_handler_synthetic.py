#!/usr/bin/env python3
"""Synthetic tests for docx_handler (session 18).

Creates in-memory .docx fixtures using python-docx and runs them through
docx_handler.extract_from_path() with mocked vision clients. No Drive,
no Gemini, no Chroma writes.

Mirrors test_google_doc_handler_synthetic.py in structure and coverage.
"""

import io
import sys
import os
import tempfile
from pathlib import Path
from dataclasses import dataclass
from typing import Optional

# Ensure repo root is on sys.path
REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO))

from ingester.loaders.types import (
    ExtractResult,
    PAGE_MARKER_RE,
    chunk_with_locators,
    derive_locator,
    strip_markers,
)
from ingester.loaders.types.docx_handler import (
    extract_from_path,
    _walk_document,
    _serialize_table,
)


# ============================================================================
# Mock vision client (mirrors test_google_doc_handler_synthetic.py)
# ============================================================================

@dataclass
class MockOcrResult:
    sha256: str = "abc123"
    is_decorative: bool = False
    failed: bool = False
    failure_reason: str = ""
    ocr_text: str = "MOCK OCR TEXT"
    vision_input_tokens: int = 100
    vision_output_tokens: int = 50


class MockLedger:
    def __init__(self):
        self.images_seen = 0
        self.images_ocr_called = 0
        self.vision_cost_usd = 0.0


class MockVisionClient:
    def __init__(self, decorative=False, fail=False):
        self.ledger = MockLedger()
        self._decorative = decorative
        self._fail = fail
        self.calls: list[dict] = []

    def ocr_image(self, img_bytes, mime_type, use_cache=True):
        self.ledger.images_seen += 1
        self.ledger.images_ocr_called += 1
        self.ledger.vision_cost_usd += 0.001
        self.calls.append({
            "byte_size": len(img_bytes),
            "mime_type": mime_type,
        })
        return MockOcrResult(
            is_decorative=self._decorative,
            failed=self._fail,
            failure_reason="mock failure" if self._fail else "",
            ocr_text="REIMAGINED HEALTH LOGO" if not self._fail else "",
        )


# ============================================================================
# Synthetic docx fixture builders
# ============================================================================

def _build_simple_docx() -> Path:
    """Build a docx with headings, paragraphs, a table, and an inline image.
    Contains the former-collaborator name to verify scrub survival."""
    from docx import Document
    from docx.shared import Inches
    import struct

    doc = Document()

    # Heading 1
    doc.add_heading("Fertility Supplement Guide", level=1)
    doc.add_paragraph(
        "This guide by Dr. Nashat Latib covers the essential supplements "
        "for optimizing fertility outcomes in women over 35."
    )

    # Heading 2 with former-collaborator name (scrub target)
    doc.add_heading("Vitamin D Protocol", level=2)
    doc.add_paragraph(
        "Dr. Christina Massinople recommends 5000 IU of Vitamin D3 daily. "
        "This protocol has shown significant improvements in AMH levels."
    )

    # Table: supplement schedule
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Supplement", "Dose", "Timing"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows_data = [
        ["Vitamin D3", "5000 IU", "Morning with food"],
        ["CoQ10", "600 mg", "Split AM/PM"],
        ["Fish Oil", "2000 mg EPA+DHA", "With meals"],
    ]
    for r, row_data in enumerate(rows_data):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    # Another heading
    doc.add_heading("Lab Monitoring", level=2)
    doc.add_paragraph(
        "Monitor Vitamin D levels every 3 months. Target range is 50-80 ng/mL. "
        "Dr. Christina also suggests checking TSH and free T4."
    )

    # Add a minimal inline image (1x1 red PNG)
    # This is a valid 1x1 red PNG file
    png_bytes = _make_tiny_png()
    img_stream = io.BytesIO(png_bytes)
    doc.add_paragraph("Company logo:")
    doc.add_picture(img_stream, width=Inches(1.0))

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(
        "Following these protocols consistently for 90 days typically "
        "results in measurable improvements across key fertility markers."
    )

    # Save to temp file
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return Path(tmp.name)


def _make_tiny_png() -> bytes:
    """Generate a minimal valid 1x1 red PNG in pure Python."""
    import struct, zlib

    def _chunk(chunk_type, data):
        c = chunk_type + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr = _chunk(b"IHDR", ihdr_data)
    raw_row = b"\x00\xff\x00\x00"  # filter=None, R=255, G=0, B=0
    idat = _chunk(b"IDAT", zlib.compress(raw_row))
    iend = _chunk(b"IEND", b"")
    return signature + ihdr + idat + iend


def _build_no_headings_docx() -> Path:
    """Build a docx with NO heading styles — only plain paragraphs."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("This document has no headings at all.")
    doc.add_paragraph(
        "It only contains plain paragraphs to test the empty-locator "
        "fallback path."
    )
    doc.add_paragraph("Third paragraph with some fertility content about AMH.")

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return Path(tmp.name)


def _build_table_only_docx() -> Path:
    """Build a docx with only a table and no headings or images."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Supplement reference table:")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Supplement"
    table.rows[0].cells[1].text = "Benefit"
    table.rows[1].cells[0].text = "CoQ10"
    table.rows[1].cells[1].text = "Egg quality"
    table.rows[2].cells[0].text = "NAC"
    table.rows[2].cells[1].text = "Antioxidant support"
    doc.add_paragraph("End of table.")

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return Path(tmp.name)


# ============================================================================
# Tests
# ============================================================================

results = []

def run_test(name, fn):
    try:
        fn()
        results.append(("PASS", name))
        print(f"  [{' PASS '}] {name}")
    except Exception as e:
        results.append(("FAIL", name))
        print(f"  [{'*FAIL*'}] {name}")
        import traceback
        traceback.print_exc()


# ---------- Test 1: imports + public interface ----------
def test_imports():
    assert callable(extract_from_path), "extract_from_path must be callable"
    from ingester.loaders.types.docx_handler import extract
    assert callable(extract), "extract must be callable"

run_test("imports_and_public_interface", test_imports)


# ---------- Test 2: walk_document produces correct stream ----------
def test_walk_document():
    from docx import Document
    path = _build_simple_docx()
    try:
        doc = Document(str(path))
        stream = _walk_document(doc)

        kinds = [e["kind"] for e in stream]
        # Expect: heading, text (heading body), text (para),
        #         heading, text (heading body), text (para),
        #         table,
        #         heading, text (heading body), text (para + Dr. Christina),
        #         text (logo caption), image,
        #         heading, text (heading body), text (para)
        assert "heading" in kinds, f"expected heading in stream, got {kinds}"
        assert "text" in kinds, f"expected text in stream"
        assert "table" in kinds, f"expected table in stream, got {kinds}"
        assert "image" in kinds, f"expected image in stream, got {kinds}"

        headings = [e for e in stream if e["kind"] == "heading"]
        assert len(headings) == 4, f"expected 4 headings, got {len(headings)}"
        assert headings[0]["text"] == "Fertility Supplement Guide"
        assert headings[0]["level"] == 1
        assert headings[1]["level"] == 2
    finally:
        os.unlink(path)

run_test("walk_document_stream_shape", test_walk_document)


# ---------- Test 3: section markers in stitched text ----------
def test_section_markers():
    path = _build_simple_docx()
    try:
        vc = MockVisionClient()
        result = extract_from_path(path, vision_client=vc)
        markers = PAGE_MARKER_RE.findall(result.stitched_text)
        section_markers = [(u, v) for u, v in markers if u == "SECTION"]
        assert len(section_markers) == 4, (
            f"expected 4 SECTION markers, got {len(section_markers)}: {section_markers}"
        )
        # Markers should be numbered 1..4
        nums = [int(v) for _, v in section_markers]
        assert nums == [1, 2, 3, 4], f"expected [1,2,3,4], got {nums}"
    finally:
        os.unlink(path)

run_test("section_markers_in_stitched_text", test_section_markers)


# ---------- Test 4: table content in stitched text ----------
def test_table_content():
    path = _build_simple_docx()
    try:
        vc = MockVisionClient()
        result = extract_from_path(path, vision_client=vc)
        # Table should appear as pipe-delimited text
        assert "Vitamin D3" in result.stitched_text, "table cell Vitamin D3 missing"
        assert "5000 IU" in result.stitched_text, "table cell 5000 IU missing"
        assert " | " in result.stitched_text, "pipe delimiter missing from table"
        assert "---" in result.stitched_text, "header separator missing from table"
    finally:
        os.unlink(path)

run_test("table_content_in_stitched_text", test_table_content)


# ---------- Test 5: image OCR integration ----------
def test_image_ocr():
    path = _build_simple_docx()
    try:
        vc = MockVisionClient()
        result = extract_from_path(path, vision_client=vc)
        assert "[IMAGE #1: REIMAGINED HEALTH LOGO]" in result.stitched_text, (
            "OCR text not spliced into stitched text"
        )
        assert vc.ledger.images_seen == 1, f"expected 1 image seen, got {vc.ledger.images_seen}"
        assert vc.ledger.images_ocr_called == 1
        assert result.images_seen == 1
        assert result.images_ocr_called == 1
        assert result.vision_cost_usd > 0
    finally:
        os.unlink(path)

run_test("image_ocr_integration", test_image_ocr)


# ---------- Test 6: ExtractResult shape ----------
def test_extract_result_shape():
    path = _build_simple_docx()
    try:
        vc = MockVisionClient()
        result = extract_from_path(path, vision_client=vc)
        assert isinstance(result, ExtractResult)
        assert result.extraction_method == "docx_python_docx"
        assert result.source_unit_label == "section"
        assert result.pages_total == 0  # docx has no pages
        assert result.units_total == 4  # 4 headings
        assert result.extra["section_count"] == 4
        assert result.extra["table_count"] == 1
        assert result.extra["image_count_in_stream"] == 1
        assert isinstance(result.extra["per_image_records"], list)
        assert len(result.warnings) == 0  # doc has headings
    finally:
        os.unlink(path)

run_test("extract_result_shape", test_extract_result_shape)


# ---------- Test 7: chunk_with_locators integration ----------
def test_chunk_with_locators_integration():
    path = _build_simple_docx()
    try:
        vc = MockVisionClient()
        result = extract_from_path(path, vision_client=vc)
        chunks = chunk_with_locators(result.stitched_text)
        assert len(chunks) >= 1, "expected at least 1 chunk"
        for chunk in chunks:
            # Markers must be stripped from final text
            assert "[SECTION " not in chunk["text"], (
                f"marker leaked into chunk text: {chunk['text'][:80]}"
            )
            # Locator should be populated (doc has headings)
            assert chunk["display_locator"] is not None or chunk["display_locator"] == "", (
                "locator should be set on chunks from a doc with headings"
            )
        # At least one chunk should have a § locator
        locators = [c["display_locator"] for c in chunks if c["display_locator"]]
        assert any("§" in loc for loc in locators), (
            f"expected at least one § locator, got {locators}"
        )
    finally:
        os.unlink(path)

run_test("chunk_with_locators_integration", test_chunk_with_locators_integration)


# ---------- Test 8: scrub fires on former-collaborator name ----------
def test_scrub_fires():
    path = _build_simple_docx()
    try:
        vc = MockVisionClient()
        result = extract_from_path(path, vision_client=vc)
        chunks = chunk_with_locators(result.stitched_text)
        total_replacements = sum(c["name_replacements"] for c in chunks)
        assert total_replacements >= 2, (
            f"expected at least 2 name replacements (Christina + Massinople), "
            f"got {total_replacements}"
        )
        # Verify no leaked names in final chunk text
        for chunk in chunks:
            text_lower = chunk["text"].lower()
            assert "christina" not in text_lower, (
                f"'Christina' leaked into chunk: {chunk['text'][:100]}"
            )
            assert "massinople" not in text_lower, (
                f"'Massinople' leaked into chunk: {chunk['text'][:100]}"
            )
        # Section markers must survive scrub (scrub shouldn't touch them)
        # This is already verified by test 7 (locators populated), but
        # let's also check the raw stitched text still has markers
        assert "[SECTION 1]" in result.stitched_text, (
            "scrub must not touch [SECTION N] markers"
        )
    finally:
        os.unlink(path)

run_test("scrub_fires_on_former_collaborator", test_scrub_fires)


# ---------- Test 9: no-headings fallback ----------
def test_no_headings():
    path = _build_no_headings_docx()
    try:
        result = extract_from_path(path, vision_client=None)
        assert result.units_total == 0, f"expected 0 headings, got {result.units_total}"
        assert result.extra["section_count"] == 0
        assert len(result.warnings) == 1, "expected 1 warning about no headings"
        assert "no heading styles found" in result.warnings[0]
        chunks = chunk_with_locators(result.stitched_text)
        assert len(chunks) >= 1
        # All locators should be None (no markers to derive from)
        for chunk in chunks:
            assert chunk["display_locator"] is None, (
                f"expected None locator, got {chunk['display_locator']}"
            )
    finally:
        os.unlink(path)

run_test("no_headings_fallback", test_no_headings)


# ---------- Test 10: table-only document ----------
def test_table_only():
    path = _build_table_only_docx()
    try:
        result = extract_from_path(path, vision_client=None)
        assert "CoQ10" in result.stitched_text
        assert "Egg quality" in result.stitched_text
        assert result.extra["table_count"] == 1
        assert result.extraction_method == "docx_python_docx"
    finally:
        os.unlink(path)

run_test("table_only_document", test_table_only)


# ---------- Test 11: decorative image is dropped ----------
def test_decorative_image_dropped():
    path = _build_simple_docx()
    try:
        vc = MockVisionClient(decorative=True)
        result = extract_from_path(path, vision_client=vc)
        # Decorative images should NOT appear in stitched text
        assert "[IMAGE #" not in result.stitched_text, (
            "decorative image should not appear in stitched text"
        )
        assert vc.ledger.images_seen == 1  # still counted
    finally:
        os.unlink(path)

run_test("decorative_image_dropped", test_decorative_image_dropped)


# ---------- Test 12: serialize_table helper ----------
def test_serialize_table():
    rows = [["A", "B", "C"], ["1", "2", "3"], ["x", "y", "z"]]
    result = _serialize_table(rows)
    lines = result.split("\n")
    assert lines[0] == "A | B | C", f"header wrong: {lines[0]}"
    assert lines[1] == "--- | --- | ---", f"separator wrong: {lines[1]}"
    assert lines[2] == "1 | 2 | 3"

    # Single row — no separator
    single = _serialize_table([["only", "row"]])
    assert "---" not in single

    # Empty
    assert _serialize_table([]) == ""

run_test("serialize_table_helper", test_serialize_table)


# ============================================================================
# Summary
# ============================================================================

print(f"\n{'='*60}")
passing = sum(1 for r in results if r[0] == "PASS")
failing = sum(1 for r in results if r[0] == "FAIL")
print(f"  {passing}/{len(results)} passing, {failing} failing")
if failing:
    print("  FAILED:")
    for status, name in results:
        if status == "FAIL":
            print(f"    - {name}")
    sys.exit(1)
